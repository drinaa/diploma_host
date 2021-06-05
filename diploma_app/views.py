from django.shortcuts import render, redirect, reverse
from .forms import LoginForm
from django.contrib import auth
from django.contrib.auth.models import User
from .models import pupil, diplomas, points
from django.db.models import Sum
from datetime import date
import docx
import io
from django.http import HttpResponse
import os
# from fpdf import FPDF


def index(request):
    return render(request, "login.html")


def login_page(request):
    if request.method == "POST":
        login = request.POST["username"]
        password = request.POST["password"]
        user = auth.authenticate(username=login, password=password)
        if user is not None:
            # Правильный пароль и пользователь "активен"
            auth.login(request, user)
            # Перенаправление на "правильную" страницу
            if user.is_staff:
                template = "../teacher"
            else:
                p_user = User.objects.get(username=login)
                request.session['pupil_id'] = pupil.objects.get(pupil_id=p_user.id).id
                request.session['cur_grade'] = pupil.objects.get(pupil_id=p_user.id).grade
                request.session['grade'] = pupil.objects.get(pupil_id=p_user.id).grade
                request.session['nomin'] = ""
                template = "../pupil"
            return redirect(template)
        else:
            context = {'error': 'Введён неверный пароль'}
            return render(request, "login.html", context)
    else:
        context = {'form': LoginForm()}
        return render(request, "login.html", context)

def base(request):
    return render(request, 'base.html')

def pupil_page(request):
    if not request.user.is_authenticated:
        return redirect("index")
    if request.user.is_authenticated and request.user.is_staff:
        return redirect("../teacher")
    return render(request, 'pupil.html')

def teacher_page(request):
    if not request.user.is_authenticated:
        return redirect("index")
    if request.user.is_authenticated and not request.user.is_staff:
        return redirect("../pupil")

    if 'show_portfolio_teach' in request.POST:
        teach_nomin = request.POST["teach_nomin"]
        teach_grade = request.POST["teach_grade"]
        teach_year = request.POST["teach_year"]
        data = diplomas.objects.select_related("pupil").filter(nomination__contains=teach_nomin,
                                                               grade__contains=teach_grade,
                                                               year__contains=teach_year).values('nomination', 'year',
                                                                                                 'grade',
                                                                                                 'pupil_id').annotate(
            points=Sum('point'))
        for elem in data:
            elem["last_name"] = pupil.objects.get(id=elem["pupil_id"]).last_name
            elem["first_name"] = pupil.objects.get(id=elem["pupil_id"]).first_name
        return render(request, 'teacher.html', {"portfolios": data})

    return render(request, 'teacher.html')

def show_grade(request):
    request.session['nomin'] = ""
    if request.GET.get("grade", "")!="":
        request.session['grade'] = request.GET.get("grade", "")
    diploma = diplomas.objects.filter(pupil=request.session["pupil_id"], grade__contains=request.session['grade'], nomination__contains=request.session['nomin'])
    return render(request, 'pupil.html', {"diplomas" : diploma, "grade" : request.session['grade']})

def show_nomin(request):
    request.session['nomin'] = request.POST["nomination"]
    diploma = diplomas.objects.filter(pupil=request.session["pupil_id"], nomination__contains=request.session['nomin'],
        grade__contains=request.session['grade'])
    sum_point = diplomas.objects.filter(pupil=request.session["pupil_id"], nomination__contains=request.session['nomin'],
        grade__contains=request.session['grade']).values('nomination','year', 'grade', 'pupil_id').annotate(points=Sum('point'))
    if not sum_point:
        sum_points = 0
    else:
        for elem in sum_point:
            sum_points = elem["points"]
    return render(request, 'pupil.html', {"diplomas" : diploma, "grade" : request.session['grade'], "nomination" : request.session['nomin'], "sum_points" : sum_points})

def add_diploma(request):
    pupil_id = request.session["pupil_id"]
    grade = request.session["cur_grade"]
    cur_year = date.today().year
    cur_month = date.today().month
    if cur_month > 4:
        cur_year += 1
        if cur_month < 9:
            grade += 1

    name = request.POST["add_descrip"]
    level = request.POST["add_level"]
    part1 = request.POST["add_part1"]
    part2 = 0
    place = 0

    if 'project' in request.POST:
        type_dipl = '0'
        code = type_dipl+level+part1
        nomination = request.POST["add_nomin"]
    else:
        part2 = request.POST["add_part2"]
        place = request.POST["add_place"]
        if 'olimp' in request.POST:
            type_dipl = '1'
            nomination = 'интеллектуальная'
        else:
            type_dipl = '2'
            if 'sport' in request.POST:
                nomination = 'спортивная'
            else:
                nomination = 'творческая'
        code = type_dipl + level+place+part1

    add_point = points.objects.get(code=code).point

    if 'project' not in request.POST:
        if part2 == 1:
            add_point /= 2

    data = diplomas(pupil_id = pupil_id, grade = grade, year = cur_year, name = name, nomination = nomination,
                    type = type_dipl, level = level, place = place, part1 = part1, part2 = part2,
                    point = add_point)
    data.save()
    return redirect("../pupil")

def create_report(request):
    report_year = request.POST["report_year"]
    list_nomination = ['интеллектуальная', 'спортивная', 'творческая']
    doc = docx.Document()

    doc.add_heading('Список номинантов за {report_year} год'.format(report_year = report_year), 0)


    for i in range (1,12):
        par = doc.add_heading('{i} класс'.format(i=i))
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        j = 0
        for nomin in list_nomination:
            data = diplomas.objects.select_related("pupil").filter(year__contains=report_year, grade__contains=i, nomination__contains=nomin).values('year', 'nomination', 'pupil_id').annotate(points=Sum('point'))[:5]
            table.cell(j, 0).text = nomin + ' номинация'
            a = table.cell(j, 0).merge(table.cell(j,1))
            j += 1
            for elem in data:
                table.add_row()
                last_name = pupil.objects.get(id=elem["pupil_id"]).last_name
                first_name = pupil.objects.get(id=elem["pupil_id"]).first_name
                table.cell(j,0).text = last_name + ' ' + first_name
                table.cell(j,1).text = str (elem["points"])
                j += 1
    doc.save(f'Список номинантов за {report_year} год.docx'.format(report_year = report_year))
    doc_io = io.BytesIO()  # create a file-like object
    doc.save(doc_io)  # save data to file-like object
    doc_io.seek(0)  # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = f"attachment; filename=Year report.docx"

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response

    # pdf = FPDF()
    # pdf.add_page()
    # pdf.set_font("Arial", size=12)
    # pdf.cell(200, 10, txt="Welcome to Python!", ln=1, align="C")
    # pdf.output("simple_demo.pdf")

    return redirect("../teacher")

def delete_diplomas(request):
    elems_arr = request.POST.getlist('delete_row')
    for elem in elems_arr:
        diplomas.objects.get(id=elem).delete()
    return redirect("portfolio/pupil")